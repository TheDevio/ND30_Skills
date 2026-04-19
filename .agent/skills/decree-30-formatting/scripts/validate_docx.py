"""
Script kiểm tra định dạng file .docx theo Nghị định 30/2020/NĐ-CP
Yêu cầu: pip install python-docx
Sử dụng: python validate_docx.py <đường_dẫn_file.docx>
"""

import sys
import os

try:
    from docx import Document
    from docx.shared import Mm, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Cần cài đặt python-docx: pip install python-docx")
    sys.exit(1)


class Decree30Validator:
    """Kiểm tra file .docx theo Nghị định 30/2020/NĐ-CP"""

    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.errors = []
        self.warnings = []
        self.passed = []

    def _mm(self, emu):
        """Chuyển EMU sang mm"""
        if emu is None:
            return None
        return round(emu / 914400 * 25.4, 1)

    def check_page_setup(self):
        """Kiểm tra thiết lập trang"""
        section = self.doc.sections[0]

        # Khổ giấy A4
        w = self._mm(section.page_width)
        h = self._mm(section.page_height)
        if w and h:
            if abs(w - 210) < 2 and abs(h - 297) < 2:
                self.passed.append(f"✅ Khổ giấy A4: {w}×{h} mm")
            else:
                self.errors.append(f"❌ Khổ giấy không đúng A4: {w}×{h} mm (cần 210×297)")

        # Lề trang
        margins = {
            "Lề trên": (self._mm(section.top_margin), 20, 25),
            "Lề dưới": (self._mm(section.bottom_margin), 20, 25),
            "Lề trái": (self._mm(section.left_margin), 30, 35),
            "Lề phải": (self._mm(section.right_margin), 15, 20),
        }

        for name, (val, min_v, max_v) in margins.items():
            if val is None:
                self.warnings.append(f"⚠️ {name}: không xác định được")
            elif min_v <= val <= max_v:
                self.passed.append(f"✅ {name}: {val} mm (chuẩn: {min_v}–{max_v})")
            else:
                self.errors.append(f"❌ {name}: {val} mm (cần {min_v}–{max_v} mm)")

    def check_fonts(self):
        """Kiểm tra font chữ"""
        non_tnr = set()
        total_runs = 0
        tnr_runs = 0

        for para in self.doc.paragraphs:
            for run in para.runs:
                total_runs += 1
                font_name = run.font.name
                if font_name:
                    if font_name == "Times New Roman":
                        tnr_runs += 1
                    else:
                        non_tnr.add(font_name)
                else:
                    tnr_runs += 1  # None = kế thừa style (thường là TNR)

        if non_tnr:
            self.errors.append(
                f"❌ Dùng font không đúng: {', '.join(non_tnr)} "
                f"(cần Times New Roman)"
            )
        else:
            self.passed.append(f"✅ Font: Times New Roman ({tnr_runs}/{total_runs} runs)")

    def check_font_sizes(self):
        """Kiểm tra cỡ chữ"""
        sizes = set()
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    size_pt = run.font.size.pt
                    sizes.add(size_pt)

        valid_sizes = {11, 12, 13, 14}
        invalid = sizes - valid_sizes - {10, 10.5, 11.5, 15, 16}  # cho phép một số ngoại lệ

        if sizes:
            self.passed.append(f"✅ Cỡ chữ sử dụng: {sorted(sizes)}")
        if invalid:
            self.warnings.append(f"⚠️ Cỡ chữ bất thường: {sorted(invalid)}")

    def check_alignment(self):
        """Kiểm tra căn lề nội dung"""
        justified_count = 0
        total_content = 0

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            total_content += 1
            if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                justified_count += 1

        if total_content > 0:
            pct = round(justified_count / total_content * 100, 1)
            if pct >= 50:
                self.passed.append(
                    f"✅ Căn đều hai bên: {justified_count}/{total_content} đoạn ({pct}%)"
                )
            else:
                self.warnings.append(
                    f"⚠️ Chỉ {justified_count}/{total_content} đoạn căn đều ({pct}%)"
                )

    def check_end_marker(self):
        """Kiểm tra dấu kết thúc văn bản (./)"""
        all_text = "\n".join(p.text for p in self.doc.paragraphs)
        if "./" in all_text:
            self.passed.append("✅ Có dấu kết thúc văn bản (./)")
        else:
            self.errors.append("❌ Thiếu dấu kết thúc văn bản (./) ở cuối nội dung")

    def check_national_title(self):
        """Kiểm tra Quốc hiệu và Tiêu ngữ"""
        all_text = "\n".join(p.text for p in self.doc.paragraphs)

        if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in all_text.upper():
            self.passed.append("✅ Có Quốc hiệu")
        else:
            self.errors.append("❌ Thiếu Quốc hiệu")

        if "Độc lập" in all_text and "Tự do" in all_text and "Hạnh phúc" in all_text:
            self.passed.append("✅ Có Tiêu ngữ (Độc lập - Tự do - Hạnh phúc)")
        else:
            self.errors.append("❌ Thiếu Tiêu ngữ")

    def run_all(self):
        """Chạy tất cả kiểm tra"""
        print(f"\n{'='*60}")
        print(f"  KIỂM TRA VĂN BẢN THEO NGHỊ ĐỊNH 30/2020/NĐ-CP")
        print(f"  File: {os.path.basename(self.filepath)}")
        print(f"{'='*60}\n")

        self.check_page_setup()
        self.check_fonts()
        self.check_font_sizes()
        self.check_alignment()
        self.check_end_marker()
        self.check_national_title()

        # Kết quả
        if self.passed:
            print("── ĐẠT ──")
            for item in self.passed:
                print(f"  {item}")

        if self.warnings:
            print("\n── CẢNH BÁO ──")
            for item in self.warnings:
                print(f"  {item}")

        if self.errors:
            print("\n── LỖI ──")
            for item in self.errors:
                print(f"  {item}")

        # Tổng kết
        print(f"\n{'─'*60}")
        total = len(self.passed) + len(self.warnings) + len(self.errors)
        print(f"  Tổng: {total} mục kiểm tra")
        print(f"  ✅ Đạt: {len(self.passed)}")
        print(f"  ⚠️  Cảnh báo: {len(self.warnings)}")
        print(f"  ❌ Lỗi: {len(self.errors)}")

        if not self.errors:
            print("\n  🎉 VĂN BẢN ĐẠT CHUẨN NGHỊ ĐỊNH 30!")
        else:
            print(f"\n  ⛔ CẦN SỬA {len(self.errors)} LỖI TRƯỚC KHI BAN HÀNH")

        print(f"{'─'*60}\n")
        return len(self.errors) == 0


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Sử dụng: python validate_docx.py <file.docx>")
        print("Ví dụ:   python validate_docx.py 'quyet_dinh.docx'")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"❌ Không tìm thấy file: {filepath}")
        sys.exit(1)

    validator = Decree30Validator(filepath)
    success = validator.run_all()
    sys.exit(0 if success else 1)
